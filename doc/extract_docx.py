# -*- coding: utf-8 -*-
"""临时脚本：从 docx 提取正文到文本"""
import os
from pathlib import Path
from docx import Document

doc_dir = Path(__file__).resolve().parent / "word"
out_path = Path(__file__).resolve().parent / "md" / "论文正文提取.txt"

# 找第一个 .docx（排除 Word 临时文件 ~$）
docx_files = [f for f in doc_dir.glob("*.docx") if not f.name.startswith("~$")]
if not docx_files:
    print("未找到 docx 文件")
    exit(1)

doc_path = docx_files[0]
print("读取:", doc_path.name)
doc = Document(str(doc_path))

lines = []
for p in doc.paragraphs:
    lines.append(p.text)
for table in doc.tables:
    for row in table.rows:
        lines.append(" | ".join(cell.text.strip() for cell in row.cells))

out_path.parent.mkdir(parents=True, exist_ok=True)
with open(out_path, "w", encoding="utf-8") as f:
    f.write("\n".join(lines))
print("已写入:", out_path)
